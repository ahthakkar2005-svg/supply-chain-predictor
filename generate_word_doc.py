import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns

def set_run_font(run, font_name='Liberation Serif', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_paragraph(doc, text, justify=True, size=12):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, size=size)
    return p

def add_data_dictionary_table(doc, dictionary_data):
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    
    headers = ['Field Name', 'Data Type', 'Description']
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row_data in dictionary_data:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            p = row_cells[i].paragraphs[0]
            run = p.add_run(cell_text)
            set_run_font(run, size=11)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(ns.qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def main():
    doc = Document()
    
    # Configure global styles explicitly as requested
    
    # Normal Text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Liberation Serif'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.line_spacing = 1.0

    # Title (Level 0)
    title_style = doc.styles['Title']
    title_style.font.name = 'Liberation Serif'
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.line_spacing = 1.0

    # Heading 1
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Liberation Serif'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(0, 0, 0)
    h1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h1_style.paragraph_format.line_spacing = 1.0

    # Heading 2
    h2_style = doc.styles['Heading 2']
    h2_style.font.name = 'Liberation Serif'
    h2_style.font.size = Pt(14)
    h2_style.font.bold = True
    h2_style.font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2_style.paragraph_format.line_spacing = 1.0

    # Heading 3
    h3_style = doc.styles['Heading 3']
    h3_style.font.name = 'Liberation Serif'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_style.paragraph_format.line_spacing = 1.0

    # Adding page number (Arabic numbers, Center aligned)
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run() 
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(ns.qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(ns.qn('xml:space'), 'preserve')
        instrText.text = "PAGE  \\* MERGEFORMAT"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(ns.qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(ns.qn('w:fldCharType'), 'end')
        
        run._r.clear_content()
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        set_run_font(run, size=12)

    add_heading(doc, "Supply Chain Disruption Predictor", level=0)
    
    # --- Content ---

    # 4. Acknowledgement
    add_heading(doc, "4. Acknowledgement", level=1)
    add_paragraph(doc, "I would like to express my profound gratitude to my guiding professor and the institution for their invaluable support, guidance, and encouragement throughout the duration of this Capstone Project. Their mentorship was instrumental in the successful completion of the 'Supply Chain Disruption Predictor' system.")
    doc.add_page_break()

    # 5. Index
    add_heading(doc, "5. Index", level=1)
    add_paragraph(doc, "Note: You may need to right-click below and select 'Update Field' -> 'Update entire table' in MS Word to render the page numbers properly.", justify=False, size=10)
    add_toc(doc)
    doc.add_page_break()

    # 6. UML Diagrams
    add_heading(doc, "6. UML Diagrams / System Diagrams", level=1)
    add_paragraph(doc, "The following architectural and design diagrams illustrate the core structure of the predictive supply chain system:")

    diagrams = [
        ("Use Case Diagram", r'c:\Users\Archi Thakkar\Downloads\use_case_diagram.png'),
        ("Class Diagram", r'c:\Users\Archi Thakkar\Downloads\class_diagram_capstone_.jpg.jpeg'),
        ("Object Diagram 1", r'c:\Users\Archi Thakkar\Downloads\object_diagram_capstone_.jpg.jpeg'),
        ("Object Diagram 2", r'c:\Users\Archi Thakkar\Downloads\object_diagram_capstone1_.jpg.jpeg'),
        ("Sequence Diagram", r'c:\Users\Archi Thakkar\Downloads\sequence_capstone.jpg.jpeg')
    ]

    for d_title, d_path in diagrams:
        add_heading(doc, d_title, level=2)
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(d_path, width=Inches(5.0))
        except Exception as e:
            add_paragraph(doc, f"[ ERROR INSERTING DIAGRAM: {e} ]", justify=False)
            
    doc.add_page_break()

    # 7. Data Dictionary
    add_heading(doc, "7. Data Dictionary / Data Sets", level=1)
    add_paragraph(doc, "The platform harnesses the following primary data schemas to store and process prediction, reporting, and news-based sentiment metrics on supply chain disruption.")
    
    add_heading(doc, "Data Dictionary", level=2)
    add_paragraph(doc, "1. NewsArticle:")
    add_data_dictionary_table(doc, [
        ['id', 'String', 'Unique document identifier.'],
        ['source', 'String', 'Publication name fetching the news.'],
        ['sentiment', 'Enum', 'Positive / Neutral / Negative.'],
        ['disruption_type', 'Enum', 'Weather, Labor, Logistics, etc.']
    ])
    add_paragraph(doc, "")

    add_paragraph(doc, "2. Prediction / Alert:")
    add_data_dictionary_table(doc, [
        ['id', 'String', 'Alert identifier.'],
        ['risk_level', 'Enum', 'Risk severity ranging from Low to High.'],
        ['disruption_type', 'Enum', 'Categorization of disruption anomaly.']
    ])
    add_paragraph(doc, "")

    add_paragraph(doc, "3. RegionRisk:")
    add_data_dictionary_table(doc, [
        ['region_name', 'String', 'Target geographical location.'],
        ['risk_level', 'Enum', 'Computed AI risk severity (Low, Medium, High).'],
        ['alerts', 'List', 'Embedded collection of current regional warnings.']
    ])
    doc.add_page_break()

    # 8. Screen Layouts
    add_heading(doc, "8. Screen Layouts", level=1)
    add_paragraph(doc, "The software application delivers an intuitive user interface for maximum situational awareness:")
    
    add_heading(doc, "Dashboard Layout", level=2)
    add_paragraph(doc, "Top Panel: Executive Summary KPIs (Active Alerts, Regions Monitored, Critical Disruptions).")
    try:
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.add_run().add_picture('dashboard.png', width=Inches(5.5))
    except Exception:
        add_paragraph(doc, "\n[ INSERT DASHBOARD SCREENSHOT HERE ]\n", justify=False, size=14)
    
    add_heading(doc, "Mobile Responsive View", level=2)
    add_paragraph(doc, "Stack-based representation prioritizing KPI alerts on top.")
    try:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run().add_picture('dashboard_mobile.png', width=Inches(3.0))
    except Exception:
        add_paragraph(doc, "\n[ INSERT MOBILE VIEW SCREENSHOT HERE ]\n", justify=False, size=14)

    pages = [
        ("Predictions View", "Interactive forecasting based on What-If Scenario and AI simulations.", "predictions.png"),
        ("Alerts View", "Comprehensive list of system-generated priority alerts.", "alerts.png"),
        ("Suppliers View", "Detailed overview of active suppliers and respective risk assessments.", "suppliers.png"),
        ("Ports View", "Geospatial mapping and analytical deep dive into global port throughput.", "ports.png"),
        ("Analytics View", "Advanced macroscopic indicators and temporal disruption trending analysis.", "analytics.png"),
        ("Settings View", "Customizable platform configurations and API integrations.", "settings.png")
    ]
    
    for title, desc, img_name in pages:
        add_heading(doc, title, level=2)
        add_paragraph(doc, desc)
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_name, width=Inches(5.5))
        except Exception:
            pass

    doc.add_page_break()

    # 9. Report Layouts
    add_heading(doc, "9. Report Layouts", level=1)
    add_paragraph(doc, "The reporting module provides automated periodic overviews of historical delays and systemic risks.")
    add_heading(doc, "Weekly Threat Assessment Report", level=2)
    add_paragraph(doc, "- Title Header: Time period and target domain.\n- Section A: Statistical abstract of top forecasted hazards.\n- Section B: Tabular data representation listing specific suppliers or regions experiencing turbulence.\n- Section C: Evaluative graph charting 30-day trailing disruption probability indices vs. predicted actuals.")
    doc.add_page_break()
    
    # 10. Sample Coding
    add_heading(doc, "10. Sample Coding", level=1)
    add_paragraph(doc, "Below is a sample FastAPI core endpoint highlighting key implementations for fetching dashboard summaries:")
    add_paragraph(doc, 'from fastapi import APIRouter\nfrom app.models import DashboardSummary\n\nrouter = APIRouter()\n\n@router.get("/api/dashboard/summary", response_model=DashboardSummary)\nasync def get_dashboard_summary():\n    return {\n        "active_alerts": 12,\n        "high_risk_regions": 2,\n        "overall_status": "Moderate Risk"\n    }')
    doc.add_page_break()

    # 11. Future Enhancements
    add_heading(doc, "11. Future Enhancements", level=1)
    add_paragraph(doc, "1. API Integration with modern ERP platforms (SAP, Oracle) for direct tracking against active purchase orders.")
    add_paragraph(doc, "2. Advanced Language Models to dynamically extract hidden relationship networks influencing logistics pathways.")
    add_paragraph(doc, "3. Implementation of a mobile, natively optimized application for instant critical-threat push notifications for on-the-go managers.")
    doc.add_page_break()
    
    # 12. Conclusion
    add_heading(doc, "12. Conclusion", level=1)
    add_paragraph(doc, "The Supply Chain Disruption Predictor demonstrates how leveraging predictive models, open data, and modern web frameworks creates massive operational value. By effectively mapping and simulating geopolitical, economic, and logistical risks into one cohesive, automated risk-intelligence framework, supply professionals can respond to crises before critical bottlenecks manifest.")
    doc.add_page_break()

    # 13. Bibliography
    add_heading(doc, "13. Bibliography", level=1)
    add_paragraph(doc, "[1] React Documentation. Available at: https://reactjs.org/\n[2] FastAPI Documentation. Available at: https://fastapi.tiangolo.com/\n[3] Meta Prophet - Open Source Forecasting. Available at: https://facebook.github.io/prophet/\n[4] MongoDB Python Driver API Notes. Available at: https://pymongo.readthedocs.io/")

    # Save the document
    doc.save('Project_Documentation_V7.docx')
    print("Document saved successfully as Project_Documentation_V7.docx")

    try:
        from docxcompose.composer import Composer
        cover_path = r'c:\Users\Archi Thakkar\Downloads\iMCA_coverPage_Capstone Project_1f5a21ad28259f7daa70f1540a497983.docx'
        cert_path = r'c:\Users\Archi Thakkar\Downloads\iMCA_insti_Certificate_CapstoneProj_691c1e00df8408016cbc0ef2e10b5c27.docx'
        
        print("Merging cover page and certificate...")
        master = Document(cover_path)
        composer = Composer(master)
        
        # Merge Certificate
        composer.append(Document(cert_path))
        
        # Merge Main Output
        composer.append(Document('Project_Documentation_V7.docx'))
        
        composer.save('Final_Project_Documentation_V5.docx')
        print("Successfully merged into Final_Project_Documentation_V5.docx")
    except Exception as e:
        print(f"Merge failed: {e}")

if __name__ == '__main__':
    main()
